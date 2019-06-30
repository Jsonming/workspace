#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Time    : 2019/6/5 18:03
# @Author  : yangmingming
# @Site    : 
# @File    : word_count.py
# @Software: PyCharm

import os
import sys
from docx import Document


def table_to_cell(tables):
    cell_list = [c.text for t in tables for r in t.rows for c in r.cells]
    cells_new = list(set(cell_list))
    cells = sorted(cells_new, key=cell_list.index)
    return cells


def article_count(file):
    document = Document(file)
    article, article_length = '', 0
    for paragraph in document.paragraphs:
        article_length += 1
        paragraph_text = paragraph.text.strip()
        article += paragraph_text

    cells = []
    cells.extend(table_to_cell(document.tables))
    article_length += len(table_to_cell(document.tables))
    cell_list = [c.text for t in document.tables for r in t.rows for c in r.cells]
    cells_new = list(set(cell_list))
    cells = sorted(cells_new, key=cell_list.index)

    cell_tables = [c.tables for t in document.tables for r in t.rows for c in r.cells if c.tables]
    for cell_table in cell_tables:
        cells.extend(table_to_cell(cell_table))
    article += ''.join(list(set(cells)))

    article = article.replace('\n', '').replace('\t', '').replace('\r', '')
    article = ''.join(article.split())
    article_length += len(article)

    return article_length


def list_file(folder):
    """
        get all file
    :param folder:
    :return:
    """
    file_list = []
    files = os.listdir(folder)
    for file in files:
        file_name = os.path.join(folder + "\\" + file)
        if os.path.isdir(file_name):
            file_list.extend(list_file(file_name))
        else:
            file_list.append(file_name)
    return file_list


def save_txt(result, file):
    """ 保存 text"""
    with open(file, 'a', encoding='utf8')as f:
        f.write('\n'.join(result))


folder = sys.argv[1]
result = []
for file_name in list_file(folder):
    file = os.path.join(folder, file_name)
    article_length = article_count(file)
    print(file_name, article_length)
    result.append('{} {}'.format(file_name, article_length))
save_txt(result, 'word_count.txt')



